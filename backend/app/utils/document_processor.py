"""
Document processing utilities for handling PDF, DOCX, and TXT files
"""
import os
from pathlib import Path
from typing import Tuple, List
import PyPDF2
from docx import Document as DocxDocument
import re


class DocumentProcessor:
    """Handles document parsing and text extraction"""
    
    ALLOWED_EXTENSIONS = {'.pdf', '.txt', '.docx'}
    
    @staticmethod
    def is_valid_file(filename: str) -> bool:
        """Check if file has allowed extension"""
        return Path(filename).suffix.lower() in DocumentProcessor.ALLOWED_EXTENSIONS
    
    @staticmethod
    def extract_text_from_pdf(file_path: str) -> str:
        """Extract text from PDF file"""
        text = ""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text += page.extract_text()
        except Exception as e:
            raise ValueError(f"Error reading PDF file: {str(e)}")
        return text
    
    @staticmethod
    def extract_text_from_docx(file_path: str) -> str:
        """Extract text from DOCX file"""
        text = ""
        try:
            doc = DocxDocument(file_path)
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + "\n"
        except Exception as e:
            raise ValueError(f"Error reading DOCX file: {str(e)}")
        return text
    
    @staticmethod
    def extract_text_from_txt(file_path: str) -> str:
        """Extract text from TXT file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
        except UnicodeDecodeError:
            # Try with different encoding
            with open(file_path, 'r', encoding='latin-1') as file:
                text = file.read()
        except Exception as e:
            raise ValueError(f"Error reading TXT file: {str(e)}")
        return text
    
    @staticmethod
    def extract_text(file_path: str) -> str:
        """Extract text from document based on file type"""
        file_ext = Path(file_path).suffix.lower()
        
        if file_ext == '.pdf':
            return DocumentProcessor.extract_text_from_pdf(file_path)
        elif file_ext == '.docx':
            return DocumentProcessor.extract_text_from_docx(file_path)
        elif file_ext == '.txt':
            return DocumentProcessor.extract_text_from_txt(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_ext}")
    
    @staticmethod
    def clean_text(text: str) -> str:
        """Clean and normalize text"""
        # Remove extra whitespace
        text = re.sub(r'\s+', ' ', text)
        # Remove special characters but keep punctuation
        text = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]', '', text)
        return text.strip()
    
    @staticmethod
    def split_into_sentences(text: str) -> List[str]:
        """Split text into sentences"""
        # Simple sentence splitting based on common delimiters
        sentences = re.split(r'(?<=[.!?])\s+', text)
        # Filter out empty sentences
        sentences = [s.strip() for s in sentences if s.strip()]
        return sentences
    
    @staticmethod
    def split_into_passages(text: str, window_size: int = 3) -> List[Tuple[str, int, int]]:
        """Split text into passages (sentence windows) with position tracking"""
        sentences = DocumentProcessor.split_into_sentences(text)
        passages = []
        
        for i in range(len(sentences) - window_size + 1):
            passage = " ".join(sentences[i:i + window_size])
            # Calculate approximate start and end positions
            start_pos = sum(len(s) + 1 for s in sentences[:i])
            end_pos = start_pos + len(passage)
            passages.append((passage, start_pos, end_pos))
        
        return passages
